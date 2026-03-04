from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def txt_to_word(input_txt_path: str, output_docx_path: str) -> None:
    """
    Convert a formatted TXT file into a Word (.docx) document.

    Expected TXT structure:
    - Projects are separated by a blank line (two newlines: '\n\n').
    - Each project block starts with a title line (often wrapped with **...**).
    - The next line is a single line containing fields separated by '|',
      e.g. "Role: ... | Duration: ... | ID: ..."
    - Subsequent lines may contain additional fields separated by '|',
      e.g. "Sponsor: ... | Awarded Amount: ... | Parties: ... | Key Tech: ..."

    Output formatting:
    - Project title is bold.
    - Other lines are indented to visually nest under the title.
    - Keys before ':' are bold (e.g., "Role:", "Duration:", etc.).
    """

    doc = Document()

    # -------------------------------------------------------------------------
    # Define (or reuse) a paragraph style for body text
    # -------------------------------------------------------------------------
    # We create a custom style named 'BodyText' if it does not exist.
    # This style uses Cambria 11pt for consistent document formatting.
    if "BodyText" not in doc.styles:
        style = doc.styles.add_style("BodyText", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Cambria"
        style.font.size = Pt(11)
    else:
        style = doc.styles["BodyText"]

    # -------------------------------------------------------------------------
    # Read the input TXT content
    # -------------------------------------------------------------------------
    with open(input_txt_path, "r", encoding="utf-8") as file:
        content = file.read()

    # Split projects by blank lines
    projects = content.strip().split("\n\n")

    for project in projects:
        lines = project.strip().split("\n")
        if not lines:
            continue

        # ---------------------------------------------------------------------
        # Title paragraph (bold)
        # ---------------------------------------------------------------------
        # The title line may contain Markdown-like asterisks (e.g., "**Title**").
        # We strip '*' and surrounding whitespace to keep a clean title.
        title = lines[0].strip(" *")
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(11)

        # Space before the title to separate projects visually
        p_title.paragraph_format.space_before = Pt(8)

        # ---------------------------------------------------------------------
        # Body lines (indented, with bold keys)
        # ---------------------------------------------------------------------
        # The first non-title line is treated as a single combined row (Role/Duration/ID)
        # separated by ' | '. Subsequent lines are split into separate paragraphs.
        for i, line in enumerate(lines[1:]):
            line = line.strip()
            if not line:
                continue

            # First details line: keep all segments on one line in Word
            if i == 0:
                p = doc.add_paragraph(style="BodyText")
                p.paragraph_format.left_indent = Pt(16)  # Indent to nest under title
                p.paragraph_format.space_after = Pt(1)

                segments = [seg.strip() for seg in line.split("|")]
                for idx, seg in enumerate(segments):
                    if ":" in seg:
                        key, value = seg.split(":", 1)
                        run = p.add_run(f"{key.strip()}: ")
                        run.bold = True
                        p.add_run(value.strip())
                    else:
                        p.add_run(seg.strip())

                    # Add a visual separator between segments
                    if idx != len(segments) - 1:
                        p.add_run("  |  ")

            # Additional lines: split by '|' and write each segment as its own paragraph
            else:
                sublines = [s.strip() for s in line.split("|")]
                for sub in sublines:
                    if not sub:
                        continue

                    p = doc.add_paragraph(style="BodyText")
                    p.paragraph_format.left_indent = Pt(16)
                    p.paragraph_format.space_after = Pt(1)

                    if ":" in sub:
                        key, value = sub.split(":", 1)
                        run = p.add_run(f"{key.strip()}: ")
                        run.bold = True
                        p.add_run(value.strip())
                    else:
                        p.add_run(sub.strip())

    # -------------------------------------------------------------------------
    # Save the Word document
    # -------------------------------------------------------------------------
    doc.save(output_docx_path)
    print(f"✅ Successfully generated: {output_docx_path}")


# ---------------------------------------------------------------------------
# Example usage
# ---------------------------------------------------------------------------
# IMPORTANT CHANGE REQUESTED:
# - Replace the ITS-related file names with AI-related file names.
# - Keep everything else the same.
txt_to_word("AI_Projects.txt", "AI_Projects.docx")
txt_to_word("IoT_Projects.txt", "IoT_Projects.docx")